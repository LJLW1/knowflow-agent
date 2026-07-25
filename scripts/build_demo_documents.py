"""Generate legal, repository-owned DOCX/PDF parser fixtures from demo text."""

from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).parents[1]
OUTPUT = ROOT / "demo_corpus" / "generated"


def build_docx() -> None:
    document = Document()
    document.add_heading("Atlas 接入指南", level=1)
    document.add_heading("开发环境", level=2)
    document.add_paragraph("Python 版本固定为 3.11，依赖通过 uv 安装。")
    document.add_heading("启动顺序", level=2)
    document.add_paragraph("先启动 API，再启动 Streamlit 演示页。")
    document.save(OUTPUT / "onboarding.docx")


def build_pdf() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text(
        (72, 72),
        "Atlas Continuity Plan\nRTO: 30 minutes\nRPO: 24 hours\n"
        "Run a recovery drill every month.",
    )
    document.save(OUTPUT / "continuity-plan.pdf")
    document.close()


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()


if __name__ == "__main__":
    main()
