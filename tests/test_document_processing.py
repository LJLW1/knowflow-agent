from io import BytesIO

import fitz
import pytest
from docx import Document

from knowflow.document.parsers import DocumentParser, InvalidDocumentError
from knowflow.document.security import validate_filename
from knowflow.document.splitter import StructureAwareSplitter


def test_markdown_preserves_heading_path() -> None:
    parsed = DocumentParser().parse_bytes(
        "architecture.md",
        b"# Architecture\n\nOverview.\n\n## Storage\n\nSQLite stores metadata.",
    )
    assert parsed.sections[-1].heading_path == ["Architecture", "Storage"]
    assert "SQLite" in parsed.sections[-1].text


def test_text_parser_cleans_control_characters() -> None:
    parsed = DocumentParser().parse_bytes("runbook.txt", b"Deploy\x00 safely.\r\n\r\nCheck logs.")
    assert "\x00" not in parsed.full_text
    assert "Deploy safely." in parsed.full_text


def test_docx_parser_reads_heading_and_body() -> None:
    document = Document()
    document.add_heading("API", level=1)
    document.add_paragraph("POST /api/v1/query returns citations.")
    buffer = BytesIO()
    document.save(buffer)

    parsed = DocumentParser().parse_bytes("api.docx", buffer.getvalue())
    assert parsed.sections[0].heading_path == ["API"]
    assert "citations" in parsed.sections[0].text


def test_pdf_parser_records_one_based_page_number() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Incident response starts with containment.")
    data = document.tobytes()
    document.close()

    parsed = DocumentParser().parse_bytes("incident.pdf", data)
    assert parsed.sections[0].page_start == 1
    assert "containment" in parsed.sections[0].text


def test_corrupt_document_returns_stable_error() -> None:
    with pytest.raises(InvalidDocumentError, match="DOCUMENT_PARSE_FAILED"):
        DocumentParser().parse_bytes("broken.pdf", b"not-a-pdf")


@pytest.mark.parametrize("filename", ["../secret.txt", "/tmp/secret.txt", "folder/file.md"])
def test_upload_filename_rejects_path_traversal(filename: str) -> None:
    with pytest.raises(ValueError, match="INVALID_FILENAME"):
        validate_filename(filename)


def test_structure_splitter_keeps_source_location() -> None:
    parsed = DocumentParser().parse_bytes(
        "runbook.md",
        ("# Release\n\n" + "check tests and logs. " * 40).encode(),
    )
    chunks = StructureAwareSplitter(chunk_size=120, overlap=20).split(
        parsed,
        project_id="p1",
        document_id="d1",
        index_version_id="v1",
    )
    assert len(chunks) > 1
    assert all(chunk.project_id == "p1" for chunk in chunks)
    assert all(chunk.section_path == ["Release"] for chunk in chunks)

