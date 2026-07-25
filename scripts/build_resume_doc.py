"""Build the project-experience Word artifact from verified repository facts."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parents[1]
OUTPUT = ROOT / "output" / "doc" / "KnowFlow-Agent-项目经历专项简历材料.docx"
METRICS_PATH = ROOT / "reports" / "evaluation" / "hash-ci.json"

NAVY = "16324F"
BLUE = "176B87"
PALE = "EAF4F7"
GRAY = "5B6573"
WHITE = "FFFFFF"
DOCUMENT_FONT = "Arial Unicode MS"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(
    cell,
    top: int = 80,
    start: int = 100,
    bottom: int = 80,
    end: int = 100,
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def keep_row(row) -> None:
    for cell in row.cells:
        properties = cell._tc.get_or_add_tcPr()
        properties.append(OxmlElement("w:cantSplit"))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), DOCUMENT_FONT)
    fonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    fonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    fonts.set(qn("w:cs"), DOCUMENT_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([fonts, color, underline])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_title(document: Document, kicker: str, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(kicker.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Title"]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    paragraph = document.add_paragraph(subtitle)
    paragraph.style = document.styles["Subtitle"]


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_bullet(document: Document, text: str, *, style: str = "List Bullet") -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def add_label_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.add_run(text)


def add_status_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.6)
    header = table.rows[0].cells
    header[0].text = "项目"
    header[1].text = "实际结果 / 状态"
    for cell in header:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.bold = True
    for index, (name, value) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = value
        if index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, PALE)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        keep_row(table.rows[-1])


def add_qa(document: Document, number: int, question: str, answer: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(f"{number}. {question}")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    paragraph = document.add_paragraph(answer)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.space_after = Pt(4)


def apply_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = DOCUMENT_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), DOCUMENT_FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string("202A35")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(3)

    for style_name, size, color in (
        ("Title", 24, NAVY),
        ("Subtitle", 10, GRAY),
        ("Heading 1", 15, NAVY),
        ("Heading 2", 11, BLUE),
    ):
        style = document.styles[style_name]
        style.font.name = DOCUMENT_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), DOCUMENT_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(8 if style_name == "Heading 1" else 5)
        style.paragraph_format.space_after = Pt(4)

    document.styles["Title"].paragraph_format.space_after = Pt(3)
    document.styles["Subtitle"].paragraph_format.space_after = Pt(10)


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.25)
        section.left_margin = Cm(1.55)
        section.right_margin = Cm(1.55)
        section.header_distance = Cm(0.55)
        section.footer_distance = Cm(0.55)
        header = section.header.paragraphs[0]
        header.text = "KNOWFLOW AGENT  ·  项目经历专项材料"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(GRAY)
        add_page_number(section.footer.paragraphs[0])


def page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    metrics = json.loads(METRICS_PATH.read_text())
    dense = metrics["dense"]
    hybrid = metrics["hybrid"]

    document = Document()
    apply_styles(document)
    configure_sections(document)

    # Page 1
    add_title(
        document,
        "Project experience pack",
        "KnowFlow Agent",
        "基于 Hermes Agent 的企业知识库与有限工作流智能体 · 非官方独立扩展项目",
    )
    link_line = document.add_paragraph()
    add_hyperlink(link_line, "GitHub：LJLW1/knowflow-agent", "https://github.com/LJLW1/knowflow-agent")
    link_line.add_run("  ·  Python 3.11  ·  MIT")

    add_heading(document, "一句话项目介绍")
    document.add_paragraph(
        "面向团队内部文档，构建支持四类文件解析、项目隔离混合检索、"
        "带引用回答、只读 GitHub MCP 和有限子智能体流程的 Hermes Agent 扩展。"
    )

    add_heading(document, "技术栈")
    document.add_paragraph(
        "Python 3.11 / Hermes Agent 0.18.2 / FastAPI / SQLAlchemy 2.0 / Alembic / "
        "SQLite / Chroma / BGE / BM25 / RRF / Streamlit / pytest / Docker Compose / GitHub Actions"
    )

    add_heading(document, "可直接粘贴到中文简历的 4 条项目经历")
    add_bullet(
        document,
        "自研 PDF、DOCX、Markdown、TXT 文档解析与结构化切分，保留标题层级、"
        "页码和内容哈希，实现项目内去重、索引版本、同步删除和启动恢复。",
    )
    add_bullet(
        document,
        "实现本地 Embedding、BM25 与 RRF 混合检索及引用 allow-list；构建 56 条评测集，"
        f"CI Hash 基线实测 Recall@6 从 {dense['recall_at_k']:.2%} 到 "
        f"{hybrid['recall_at_k']:.2%}（仅作流程基线，不代表 BGE 质量）。",
    )
    add_bullet(
        document,
        "使用 FastAPI、Pydantic、SQLAlchemy/Alembic 统一 API、Hermes 插件和评测契约，"
        "实现项目级数据隔离、trace ID、JSON 日志、敏感字段脱敏与明确错误恢复。",
    )
    add_bullet(
        document,
        "集成 Hermes Tool Registry、Cron、delegate_task 与只读 GitHub MCP；"
        "限制目标仓库和读取工具，设计最多 3 个 child 的部分失败降级流程。"
    )

    add_heading(document, "English résumé bullets")
    add_bullet(
        document,
        "Built a four-format document pipeline with heading/page metadata, content hashing, "
        "incremental versions, crash reconciliation, and project-scoped SQLite/Chroma lifecycle.",
    )
    add_bullet(
        document,
        "Implemented dense + BM25 retrieval with RRF and citation allow-list validation; "
        "created a reproducible 56-case evaluation set and separate CI/local workflows.",
    )
    add_bullet(
        document,
        "Delivered FastAPI, Streamlit, typed contracts, structured tracing, secret redaction, "
        "Docker Compose, pytest, strict mypy, Ruff, and GitHub Actions.",
    )
    add_bullet(
        document,
        "Integrated two thin Hermes tools and a read-only GitHub MCP policy without modifying "
        "the upstream Agent Loop; bounded delegation to three child tasks with partial fallback.",
    )

    # Page 2
    page_break(document)
    add_title(document, "Scope & ownership", "STAR 与贡献边界", "把场景、工作和上游复用说清楚")
    add_heading(document, "STAR 项目介绍")
    add_label_paragraph(
        document,
        "S - Situation：",
        "企业资料分散在需求、架构、API、运维和事故文档中；普通对话演示缺少引用、隔离和可复现实验。",
    )
    add_label_paragraph(
        document,
        "T - Task：",
        "在一个月级范围内，基于 Hermes Agent 建立可公开、可测试、可部署的知识库与有限工作流 MVP。",
    )
    add_label_paragraph(
        document,
        "A - Action：",
        "按 TDD 完成领域契约、四类 Parser、结构切分、增量索引、混合检索、引用约束 RAG、"
        "FastAPI/Streamlit、Hermes 薄插件、MCP 双白名单、有限流程、日志、Docker 和 CI。",
    )
    add_label_paragraph(
        document,
        "R - Result：",
        f"形成 34 个源码模块、51 个通过的自动化测试、91.59% 覆盖率和 56 条评测集；"
        f"CI Hash 基线在 33 条可检索样本上得到 Dense Recall@6 {dense['recall_at_k']:.2%}、"
        f"Hybrid Recall@6 {hybrid['recall_at_k']:.2%}。云模型效果仍待真实 Key 测量。",
    )

    add_heading(document, "个人开发 / 集成 / 上游复用")
    add_status_table(
        document,
        [
            (
                "[自研]",
                "文档解析、清洗、结构切分、内容哈希、增量索引、项目隔离数据库、"
                "Dense+BM25+RRF、引用校验、RAG service、评测与工程交付。",
            ),
            (
                "[集成]",
                "knowledge_search、knowledge_answer 薄插件；GitHub MCP 只读配置、"
                "仓库/工具前置策略；三 child 的业务计划和部分失败语义。",
            ),
            (
                "[Hermes 复用]",
                "Agent Loop、Tool Registry、Provider、Memory、Session Search、MCP 客户端、"
                "Cron、delegate_task、Gateway、插件加载和安全审批。",
            ),
        ],
    )

    add_heading(document, "关键设计选择")
    add_bullet(document, "不引入 LangChain 重写 Agent Loop；服务逻辑由 API、插件、评测共用。")
    add_bullet(document, "TaskRun 明确为进程内任务；服务重启后标记 interrupted，不包装成持久队列。")
    add_bullet(document, "引用只能指向本轮实际送入模型的 Chunk，文档内指令一律视为不可信数据。")

    # Page 3
    page_break(document)
    add_title(
        document,
        "Evidence, not estimates",
        "真实测试与评测结果",
        "所有数字来自本机命令；未运行项保留“待测量”",
    )
    add_status_table(
        document,
        [
            ("自动化测试", "51 passed；pytest-cov 总覆盖率 91.59%。"),
            ("静态质量", "Ruff 全通过；mypy strict 检查 34 个源码文件无问题。"),
            ("评测集", "56 条，7 类；其中 33 条包含可检索的目标文档。"),
            ("演示语料", "8 个公开自建文档，覆盖 4 种格式；实际生成 25 个 Chunk。"),
            (
                "Dense CI 基线",
                f"HashEmbedding Recall@6 {dense['recall_at_k']:.2%}；"
                f"平均 {dense['mean_latency_ms']:.2f} ms，P95 {dense['p95_latency_ms']:.2f} ms。",
            ),
            (
                "Hybrid CI 基线",
                f"BM25 + Dense + RRF Recall@6 {hybrid['recall_at_k']:.2%}；"
                f"平均 {hybrid['mean_latency_ms']:.2f} ms，P95 {hybrid['p95_latency_ms']:.2f} ms。",
            ),
            (
                "索引与资源",
                f"索引 {metrics['index_time_ms']:.2f} ms；"
                f"峰值 RSS {metrics['peak_rss_mb']:.2f} MB。",
            ),
            (
                "检索来源完整性",
                f"{metrics['retrieval_source_integrity']:.0%}；只验证检索存储自洽。"
                "回答引用正确率待测量。",
            ),
            ("BGE 本机评测", "待测量：依赖已安装，权重下载受当前代理链路速度影响。"),
            (
                "云模型指标",
                "引用正确率、回答忠实度、幻觉率、工具选择准确率、任务完成率、Token/API 成本、"
                "端到端延迟、错误恢复率均待测量。",
            ),
            (
                "Docker 验收",
                "GitHub Actions Linux runner 已通过镜像构建、容器健康检查和命名卷重启持久化检查；"
                "本机 Docker Desktop 尚未启动。",
            ),
            (
                "GitHub 远端",
                "公开仓库已发布；5 个真实 Issues 对应独立分支/PR，CI 与 Docker workflow 均为绿色。",
            ),
        ],
    )
    add_heading(document, "口径说明")
    add_bullet(
        document,
        "HashEmbedding 是确定性的 CI 测试替身，适合验证评测程序和回归，"
        "不可解释为中文语义模型效果。",
    )
    add_bullet(
        document,
        "Top-K 固定为 6；评测对每条问题计算目标文档召回比例，再取宏平均。",
    )

    # Page 4
    page_break(document)
    add_title(
        document,
        "Interview narrative",
        "面试介绍",
        "一分钟版本先讲主线，三分钟版本再讲设计证据",
    )
    add_heading(document, "一分钟介绍")
    document.add_paragraph(
        "KnowFlow Agent 是我基于 Hermes Agent 0.18.2 做的企业知识库与有限工作流项目。"
        "我没有重写 Hermes 已有的 Agent Loop，而是独立实现 PDF、DOCX、Markdown、TXT "
        "解析、结构切分、SQLite 与 Chroma 项目隔离、Dense+BM25+RRF 混合检索，以及只能引用"
        "本轮检索 Chunk 的 RAG 服务。对外提供 FastAPI 和 Streamlit，并注册两个 Hermes 工具；"
        "GitHub MCP 同时限制只读工具和指定仓库。项目目前有 51 个测试、91.59% 覆盖率和 56 条评测集。"
        "由于没有云模型 Key，我只报告本地检索和工程测试数据，回答质量指标保持待测量。"
    )
    add_heading(document, "三分钟介绍")
    add_label_paragraph(
        document,
        "第 1 分钟 - 问题与范围：",
        "团队文档问答最容易被忽略的是来源引用、跨项目隔离、无答案拒答和失败恢复。"
        "我选择 Hermes 是因为它已经有循环、工具、Cron、MCP 和子智能体；"
        "我的目标是补齐企业文档垂直能力。",
    )
    add_label_paragraph(
        document,
        "第 2 分钟 - 技术实现：",
        "上传后按格式解析并保留页码/标题，内容哈希控制去重，更新产生索引版本。"
        "检索用本地 Embedding 和 Chroma，并与 BM25 通过 RRF 融合。模型只收到 Top-K Chunk，"
        "Prompt 禁止执行文档指令，返回后再用 allow-list 校验引用。"
        "API、插件和评测共享 Pydantic 契约。MCP 通过远程只读 Header、"
        "精确工具集和前置钩子三层限制；复杂任务最多三个 child，允许部分失败。",
    )
    add_label_paragraph(
        document,
        "第 3 分钟 - 验证与反思：",
        "我构建了 56 条七类问题，实际跑了检索、索引时间、延迟、资源、引用存在性、"
        "51 个 pytest、Ruff 和 strict mypy。Hash 只作为 CI 基线，BGE 和云模型数据没有跑完就不写。"
        "当前最大不足是 TaskRun 仍是进程内语义、没有 RBAC 和 OCR；下一步应先完成真实模型 badcase、"
        "校准拒答阈值和 reranker，而不是马上做 Kubernetes 或 GraphRAG。",
    )
    add_heading(document, "项目难点")
    add_bullet(document, "贡献归因：把上游 Agent 能力与个人业务/工程工作分开。")
    add_bullet(document, "引用可信：同时约束入模证据、模型输出 ID 和返回对象。")
    add_bullet(document, "项目隔离：SQLite、Chroma、任务和偏好每次访问都携带 project_id。")
    add_bullet(document, "失败语义：MCP 超时、child 部分失败和进程重启都有稳定可解释状态。")

    # Page 5
    page_break(document)
    add_title(document, "Q&A 01-08", "面试问答 (1-8)", "回答围绕实际代码和权衡，不背框架名词")
    questions = [
        (
            "为什么用 Hermes，而不是 LangChain/LangGraph？",
            "Hermes 已有稳定 Agent Loop、工具注册、MCP、Cron 和 delegate_task。"
            "本项目的问题是文档证据与工程闭环，重写循环会扩大范围；只有实验能证明新框架解决具体问题时才引入。",
        ),
        (
            "你的个人贡献是什么？",
            "四类文档流水线、增量索引、项目隔离数据层、混合检索、引用约束 RAG、"
            "FastAPI/Streamlit、Hermes 薄插件、MCP 策略、有限流程、测试、评测、Docker 和 CI。",
        ),
        (
            "如何避免把 Hermes 原功能写成你的成果？",
            "README、UPSTREAM 和 Word 都把 Agent Loop、Memory、MCP 客户端、Cron、"
            "delegate_task 等列为上游复用，并固定版本和 commit SHA。",
        ),
        (
            "为什么同时用 SQLite 和 Chroma？",
            "SQLite 管元数据、版本、偏好、任务和评测运行；Chroma 管向量检索。"
            "两边都保存 project_id 等关键标识，业务服务负责生命周期联动。",
        ),
        (
            "RRF 比直接加权分数有什么好处？",
            "BM25 与向量距离的量纲不同，直接相加需要校准。RRF 只依赖排名，"
            "实现简单且对分数尺度不敏感，适合作为一个月 MVP 的可解释融合基线。",
        ),
        (
            "引用正确是怎么保证的？",
            "模型只能看到本轮 Top-K，并返回 chunk_id；服务端把 ID 与检索结果做 allow-list 交集，"
            "Citation 再校验 chunk_id、页码、章节和证据片段。不存在的 ID 会变成拒答。",
        ),
        (
            "如何处理 Prompt Injection？",
            "系统 Prompt 明确文档是不可信数据，不执行其中指令；工具层另有仓库和工具白名单。"
            "这不是只靠一句 Prompt，而是模型、服务和工具三层约束。",
        ),
        (
            "增量索引如何工作？",
            "先计算 SHA-256；项目内相同哈希直接跳过。同名内容变化复用 document_id、"
            "先持久化 indexing 标记再发布新版本；失败时补偿，重启时按旧 Chunk 恢复。",
        ),
    ]
    for index, (question, answer) in enumerate(questions, start=1):
        add_qa(document, index, question, answer)

    # Page 6
    page_break(document)
    add_title(
        document,
        "Q&A 09-15",
        "问答复盘 (9-15)",
        "知道系统边界，通常比堆更多名词更有说服力",
    )
    questions = [
        (
            "为什么 TaskRun 不做消息队列？",
            "一个月 MVP 不需要分布式队列。我明确声明进程内语义，重启把 running 置为 interrupted，"
            "避免向面试官暗示不存在的可靠执行能力。",
        ),
        (
            "MCP 安全做了哪些限制？",
            "远程端启用 X-MCP-Readonly，Hermes 配置只 include 读取工具；"
            "插件 pre_tool_call 再校验工具名和 LJLW1/knowflow-agent 仓库，超时返回稳定错误。",
        ),
        (
            "子智能体为什么最多三个？",
            "固定为内部检索、GitHub 外部收集、引用校验，主智能体综合。"
            "上限控制成本与失败面；单 child 失败时返回 partial 并保留成功证据。",
        ),
        (
            "如何评价检索？",
            "问题记录目标文档，取 Top-K 后计算每题目标文档召回，再做宏平均；同时记录平均/P95 延迟、"
            "索引耗时和 RSS。Hash 与 BGE 报告必须分开。",
        ),
        (
            "为什么不报告幻觉率？",
            "当前没有真实云模型 Key。FakeLLM 只能验证代码路径，不代表模型质量，"
            "所以忠实度、幻觉率、工具选择与成本全部标记待测量。",
        ),
        (
            "测试覆盖了哪些失败？",
            "损坏 PDF、路径穿越、跨项目读取、重复/更新/删除、引用越界、MCP 超时、"
            "写工具/跨仓库阻断、child 部分失败、Cron 幂等、重启中断和日志脱敏。",
        ),
        (
            "下一步最值得做什么？",
            "先完成 BGE 与真实模型评测，分析 badcase，校准拒答阈值并实验 reranker；"
            "只有数据量和并发证明需要时，才升级 PostgreSQL、队列或更复杂部署。",
        ),
    ]
    for index, (question, answer) in enumerate(questions, start=9):
        add_qa(document, index, question, answer)

    add_heading(document, "项目不足")
    document.add_paragraph(
        "无 OCR/表格理解；无完整 RBAC/SSO；任务不跨进程恢复；BGE 与云模型效果尚未测量；"
        "本机 Docker Desktop 尚未启动，但 Linux runner 已通过镜像、健康检查和重启持久化验收。"
    )
    add_heading(document, "后续优化顺序")
    document.add_paragraph(
        "真实模型评测与 badcase → 拒答阈值 → reranker/元数据过滤 → 性能压测 → "
        "按需要引入 PostgreSQL 或可靠任务队列。"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
