"""Parsers for the four deliberately supported document formats."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

import fitz  # type: ignore[import-untyped]
from docx import Document

from knowflow.document.security import validate_filename
from knowflow.domain.models import MediaType


class InvalidDocumentError(ValueError):
    pass


@dataclass(slots=True)
class ParsedSection:
    text: str
    heading_path: list[str] = field(default_factory=list)
    page_start: int | None = None
    page_end: int | None = None


@dataclass(slots=True)
class ParsedDocument:
    filename: str
    media_type: MediaType
    sections: list[ParsedSection]

    @property
    def full_text(self) -> str:
        return "\n\n".join(section.text for section in self.sections)


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "".join(char for char in text if char in "\n\t" or ord(char) >= 32)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


class DocumentParser:
    _extensions = {
        ".pdf": MediaType.PDF,
        ".docx": MediaType.DOCX,
        ".md": MediaType.MARKDOWN,
        ".markdown": MediaType.MARKDOWN,
        ".txt": MediaType.TEXT,
    }

    def parse_bytes(self, filename: str, data: bytes) -> ParsedDocument:
        validate_filename(filename)
        suffix = Path(filename).suffix.lower()
        media_type = self._extensions.get(suffix)
        if media_type is None:
            raise InvalidDocumentError("UNSUPPORTED_MEDIA_TYPE")
        try:
            if media_type is MediaType.PDF:
                sections = self._parse_pdf(data)
            elif media_type is MediaType.DOCX:
                sections = self._parse_docx(data)
            elif media_type is MediaType.MARKDOWN:
                sections = self._parse_markdown(data.decode("utf-8"))
            else:
                sections = [ParsedSection(clean_text(data.decode("utf-8")))]
        except (InvalidDocumentError, UnicodeDecodeError):
            raise
        except Exception as exc:
            raise InvalidDocumentError("DOCUMENT_PARSE_FAILED") from exc
        sections = [section for section in sections if section.text.strip()]
        if not sections:
            raise InvalidDocumentError("DOCUMENT_EMPTY")
        return ParsedDocument(filename, media_type, sections)

    @staticmethod
    def _parse_pdf(data: bytes) -> list[ParsedSection]:
        try:
            document = fitz.open(stream=data, filetype="pdf")
        except Exception as exc:
            raise InvalidDocumentError("DOCUMENT_PARSE_FAILED") from exc
        try:
            return [
                ParsedSection(
                    text=clean_text(page.get_text("text")),
                    page_start=index + 1,
                    page_end=index + 1,
                )
                for index, page in enumerate(document)
            ]
        finally:
            document.close()

    @staticmethod
    def _parse_docx(data: bytes) -> list[ParsedSection]:
        document = Document(BytesIO(data))
        sections: list[ParsedSection] = []
        headings: list[str] = []
        body: list[str] = []

        def flush() -> None:
            text = clean_text("\n".join(body))
            if text:
                sections.append(ParsedSection(text=text, heading_path=headings.copy()))
            body.clear()

        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
            style = paragraph.style.name if paragraph.style else ""
            match = re.match(r"Heading (\d+)", style)
            if match:
                flush()
                level = int(match.group(1))
                headings[:] = headings[: level - 1]
                headings.append(text)
            else:
                body.append(text)
        flush()
        return sections

    @staticmethod
    def _parse_markdown(text: str) -> list[ParsedSection]:
        sections: list[ParsedSection] = []
        headings: list[str] = []
        body: list[str] = []

        def flush() -> None:
            content = clean_text("\n".join(body))
            if content:
                sections.append(ParsedSection(content, headings.copy()))
            body.clear()

        for line in text.splitlines():
            match = re.match(r"^(#{1,6})\s+(.+?)\s*#*$", line)
            if match:
                flush()
                level = len(match.group(1))
                headings[:] = headings[: level - 1]
                headings.append(clean_text(match.group(2)))
            else:
                body.append(line)
        flush()
        return sections
